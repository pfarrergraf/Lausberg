from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path
import json
import zipfile

OUT = Path('/mnt/data/antonomasie_pack')
DOCX_PATH = OUT / '03_Antonomasie_Einleger_Entwurf.docx'
MD_ANALYSE = OUT / '01_Antonomasie_Analyse.md'
JSON_PATH = OUT / '02_Antonomasie_JSON.json'
MD_EINLEGER = OUT / '04_Antonomasie_Einleger_Entwurf.md'
QUELLEN = OUT / '05_Quellen_und_Transkription.md'
README = OUT / 'README.txt'
ZIP_PATH = Path('/mnt/data/Antonomasie_Output_Paket.zip')

MAROON = '7C2130'
GOLD = '9C6B1F'
LIGHT = 'F3EFEA'
LIGHT2 = 'F7F4F1'
TEXT = '222222'
GREY = '666666'
BORDER = 'D8C9BC'

figure_json = {
  'id': 'antonomasie',
  'figure_name_de': 'Antonomasie',
  'figure_name_original': [
    {'language': 'griechisch', 'term': 'ἀντονομασία'},
    {'language': 'latein', 'term': 'antonomasia'},
    {'language': 'latein', 'term': 'pronominatio'}
  ],
  'category': 'Trope / namensersetzende Synekdoche',
  'definition_precise': 'Antonomasie ist die Ersetzung eines Eigennamens durch ein Appellativ oder eine Periphrase, die den gemeinten Referenten über eine hervorstechende Eigenschaft, Herkunft, Tat, Beziehung oder Charakterisierung identifizierbar macht.',
  'definition_simple': 'Man nennt eine Person oder Sache nicht direkt beim Namen, sondern ersetzt den Namen durch eine kennzeichnende Bezeichnung oder Umschreibung.',
  'formal_pattern': 'Eigenname -> Appellativ oder Periphrase; in der vossianischen Umkehrform: Eigenname -> Appellativ im Sinne eines Typusnamens',
  'semantic_function': 'Charakterisierung, Würdigung, Abwertung, Verdichtung, typologische Aufladung, epideiktische Hervorhebung',
  'distinguish_from': [
    'Periphrase: nur dann Antonomasie, wenn die Umschreibung den Eigennamen funktional ersetzt',
    'Epitheton: nur dann Antonomasie, wenn das Beiwort selbständig namensersetzend gebraucht wird',
    'Metapher: nicht primär Bildübertragung, sondern identifizierender Namensersatz',
    'Synekdoche allgemein: Antonomasie ist ein Spezialfall der Synekdoche für Eigennamen'
  ],
  'lausberg_reference': 'Lausberg, Handbuch der literarischen Rhetorik, §§ 580–581',
  'classical_examples': [
    {
      'original_language': 'griechisch',
      'original_text': 'ἀντονομασία ἐστὶ λέξις ἢ φράσις διὰ συνωνύμων ὀνομάτων τὸ κύριον παριστῶσα.',
      'translation_de': 'Antonomasie ist ein Ausdruck oder eine Umschreibung, die mittels synonymischer Benennungen den Eigennamen vergegenwärtigt.',
      'source': 'Tryphon, trop. p. 204,24, zitiert bei Lausberg §580',
      'notes': 'Grunddefinition; der Eigennamenbezug bleibt erhalten, obwohl der Name ersetzt wird.'
    },
    {
      'original_language': 'griechisch',
      'original_text': 'Φοῖβος',
      'translation_de': 'Phoibos',
      'source': 'Tryphon, bei Lausberg §580',
      'notes': 'Appellativischer Ersatzname für Apollo.'
    },
    {
      'original_language': 'griechisch',
      'original_text': 'Ἐννοσίγαιος',
      'translation_de': 'der Erderschütterer',
      'source': 'Tryphon, bei Lausberg §580',
      'notes': 'Appellativischer Ersatzname für Poseidon.'
    },
    {
      'original_language': 'griechisch',
      'original_text': 'γλαυκῶπις',
      'translation_de': 'die helläugige / eulenäugige',
      'source': 'Tryphon, bei Lausberg §580',
      'notes': 'Appellativischer Ersatzname für Athene; Lesung im Scan leicht unsicher, aber sachlich naheliegend.'
    },
    {
      'original_language': 'griechisch',
      'original_text': 'Λητοῦς καὶ Διὸς υἱός',
      'translation_de': 'Sohn der Leto und des Zeus',
      'source': 'Tryphon, bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Apollo.'
    },
    {
      'original_language': 'latein',
      'original_text': 'pronominatio est, quae sicuti cognomine quodam extraneo demonstrat id quod suo nomine non potest appellari',
      'translation_de': 'Pronominatio ist eine Figur, die gleichsam durch einen fremden Beinamen das bezeichnet, was mit seinem eigenen Namen nicht so benannt werden kann.',
      'source': 'Rhetorica ad Herennium 4,31,42, zitiert bei Lausberg §580',
      'notes': 'Älterer lateinischer Terminus und Funktionsbestimmung.'
    },
    {
      'original_language': 'latein',
      'original_text': 'antonomasia, quae aliquid pro nomine ponit',
      'translation_de': 'Antonomasie ist eine Figur, die etwas an die Stelle des Namens setzt.',
      'source': 'Quintilian 8,6,29–30',
      'notes': 'Knappe lateinische Grunddefinition.'
    },
    {
      'original_language': 'latein',
      'original_text': 'Tydides',
      'translation_de': 'der Sohn des Tydeus',
      'source': 'Quintilian 8,6,29–30',
      'notes': 'Ersatzname für Diomedes.'
    },
    {
      'original_language': 'latein',
      'original_text': 'Pelides',
      'translation_de': 'der Sohn des Peleus',
      'source': 'Quintilian 8,6,29–30',
      'notes': 'Ersatzname für Achill.'
    },
    {
      'original_language': 'latein',
      'original_text': 'divum pater atque hominum rex',
      'translation_de': 'Vater der Götter und König der Menschen',
      'source': 'Vergil, Aen. 1,65; bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Jupiter.'
    },
    {
      'original_language': 'latein',
      'original_text': 'eversorem Carthaginis et Numantiae',
      'translation_de': 'den Zerstörer Karthagos und Numantias',
      'source': 'Quintilian, zitiert bei Lausberg §580',
      'notes': 'Rhetorischer Ersatzname für Scipio.'
    },
    {
      'original_language': 'latein',
      'original_text': 'Romanae eloquentiae principem',
      'translation_de': 'den Fürsten der römischen Beredsamkeit',
      'source': 'Quintilian, zitiert bei Lausberg §580',
      'notes': 'Rhetorischer Ersatzname für Cicero.'
    },
    {
      'original_language': 'latein',
      'original_text': 'antonomasia est pro nomine, id est vice nominis posita',
      'translation_de': 'Antonomasie ist ein Ausdruck anstelle des Namens, das heißt an die Stelle des Namens gesetzt.',
      'source': 'Isidor 1,37,11',
      'notes': 'Spätantike Definition.'
    },
    {
      'original_language': 'latein',
      'original_text': 'Maia genitus',
      'translation_de': 'der von Maia Geborene',
      'source': 'Isidor, bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Merkur.'
    },
    {
      'original_language': 'latein',
      'original_text': 'Romanus sedendo vincit',
      'translation_de': 'Der Römer siegt durch Sitzen / Zuwarten.',
      'source': 'Varro rust. 1,2,2; bei Lausberg §580',
      'notes': 'Romanus steht antonomastisch für Q. Fabius Maximus.'
    },
    {
      'original_language': 'latein',
      'original_text': 'Troius heros',
      'translation_de': 'der trojanische Held',
      'source': 'Vergil, Aen. 6,451; bei Lausberg §580',
      'notes': 'Antonomastische Umschreibung.'
    },
    {
      'original_language': 'latein',
      'original_text': 'qui autem tradidit eum',
      'translation_de': 'der aber ihn verriet',
      'source': 'Mt 26,48; bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Judas.'
    },
    {
      'original_language': 'latein',
      'original_text': 'discipulum quem diligebat [Iesus]',
      'translation_de': 'den Jünger, den [Jesus] liebte',
      'source': 'Joh 19,26; bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Johannes; Ergänzung [Iesus] im Scan kontextuell erschlossen.'
    },
    {
      'original_language': 'französisch',
      'original_text': 'Neptune le protège, et ce dieu tutélaire ne sera pas en vain imploré.',
      'translation_de': 'Neptun schützt ihn, und dieser Schutzgott wird nicht vergeblich angerufen werden.',
      'source': 'Phèdre 2,5,621; bei Lausberg §580',
      'notes': 'Charakterisierend-steigernde Periphrase statt bloßer Namenswiederholung.'
    },
    {
      'original_language': 'französisch',
      'original_text': 'la mer qui vit tomber Icare',
      'translation_de': 'das Meer, das Ikarus stürzen sah',
      'source': 'Du Bellay, Regrets 31; bei Lausberg §580',
      'notes': 'Periphrastische Ortscharakterisierung mit mythologischer Aufladung.'
    },
    {
      'original_language': 'französisch',
      'original_text': 'cestuy-là qui conquit la toison',
      'translation_de': 'derjenige, der das Vlies gewann',
      'source': 'Du Bellay, Regrets 31; bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Jason.'
    },
    {
      'original_language': 'englisch',
      'original_text': 'the mother of mankind',
      'translation_de': 'die Mutter des Menschengeschlechts',
      'source': 'Milton, Paradise Lost 1,36; bei Lausberg §580',
      'notes': 'Periphrastischer Ersatzname für Eva.'
    },
    {
      'original_language': 'englisch',
      'original_text': 'one greater Man',
      'translation_de': 'ein größerer Mensch / der größere Mensch',
      'source': 'Milton, Paradise Lost 1,4; bei Lausberg §580',
      'notes': 'Christologische Antonomasie.'
    },
    {
      'original_language': 'latein',
      'original_text': 'quae sunt in quoque praecipua, proprii locum accipiunt, ut Fabius ... Cunctator est appellatus',
      'translation_de': 'Was an jemandem jeweils besonders hervortritt, nimmt die Stelle des Eigennamens ein; so wurde Fabius „der Zögerer“ genannt.',
      'source': 'Quintilian 8,2,11; bei Lausberg §580',
      'notes': 'Grenzfall zur normalen Bezeichnung; Beiname übernimmt Namensfunktion.'
    },
    {
      'original_language': 'griechisch',
      'original_text': 'πολλὴν Ἀφροδίτην τῷ λόγῳ περιτιθέναι',
      'translation_de': 'der Rede viel Aphrodite / viel Anmut beilegen',
      'source': 'Dion. Hal. de comp. 3; bei Lausberg §581',
      'notes': 'Beispiel der vossianischen Antonomasie: Eigenname als Typuswort für eine Eigenschaft.'
    }
  ],
  'modern_examples': [
    {
      'text': 'Die Eiserne Lady prägte ein ganzes Jahrzehnt.',
      'domain': 'Politik / Medien',
      'notes': 'Beiname ersetzt den Eigennamen wertend und identifizierend.'
    },
    {
      'text': 'Der Diktator unterschrieb den Befehl.',
      'domain': 'Geschichte / Kommentar',
      'notes': 'Nur dann saubere Antonomasie, wenn der Referent im Kontext eindeutig ist.'
    },
    {
      'text': 'Der Menschenfänger des Internets wusste genau, wie er Aufmerksamkeit bindet.',
      'domain': 'digitale Kultur',
      'notes': 'Moderne, wertende Namensersatzbildung.'
    }
  ],
  'sermon_examples': [
    {
      'text': 'Der Friedefürst betritt nicht den Palast, sondern unsere zerrissenen Häuser.',
      'purpose': 'christologische Verdichtung',
      'audience': 'Advent / Weihnachten / Gemeinde',
      'notes': 'Titel ersetzt den Namen Jesu und trägt zugleich die Auslegung.'
    },
    {
      'text': 'Der gute Hirte sucht nicht starke Tiere, sondern verlorene Schafe.',
      'purpose': 'evangelistische Zuspitzung',
      'audience': 'gemischte Gemeinde',
      'notes': 'Bekannte biblische Antonomasie mit starker Wiedererkennbarkeit.'
    },
    {
      'text': 'Der Menschenfischer zerbricht am Feuer und wird am Ufer neu gerufen.',
      'purpose': 'Petrus typologisch charakterisieren',
      'audience': 'bibelfeste Hörer',
      'notes': 'Ersetzt Petrus nicht neutral, sondern deutend.'
    }
  ],
  'effects': [
    'charakterisiert im Moment der Benennung',
    'verdichtet Tat, Würde, Herkunft oder Funktion',
    'erspart nüchterne Namenswiederholung',
    'steigert epideiktische Farbigkeit',
    'kann ehrend oder abwertend wirken',
    'öffnet typologische und homiletische Anschlüsse'
  ],
  'best_use_cases': [
    'feierliche Personenbenennung',
    'christologische Titelrede',
    'charakterisierende Würdigung oder Kritik',
    'Ersetzung klobiger Namenswiederholung',
    'typologische Verdichtung in Predigt und Rede'
  ],
  'avoid_when': [
    'der Referent aus dem Kontext nicht eindeutig identifizierbar ist',
    'nur dekorative Umschreibung ohne echten Namensersatz vorliegt',
    'der Text maximale Sachnüchternheit verlangt',
    'Metapher und Antonomasie vermischt würden, ohne dass Individualisierung gegeben ist'
  ],
  'common_confusions': [
    'bloße Periphrase',
    'bloßes Epitheton',
    'Metapher ohne eindeutigen Individualbezug',
    'allgemeiner Ehrentitel ohne Namensfunktion',
    'vossianische Antonomasie als eigene Hauptfigur statt als spätere Umkehrform'
  ],
  'negative_examples': [
    'Jesus, der Friedefürst, spricht ...',
    'der tapfere Achill',
    'ein Löwe von Mann'
  ],
  'generation_rules': [
    'Wähle eine Eigenschaft, Herkunft, Tat oder Beziehung, die den Referenten eindeutig markiert.',
    'Prüfe, ob der Eigenname wirklich ersetzt wird und nicht bloß zusätzlich genannt bleibt.',
    'Nutze möglichst kulturell oder biblisch anschlussfähige Kennzeichen.',
    'Setze die Figur dort ein, wo Benennung und Deutung zusammenfallen sollen.',
    'Behandle die vossianische Umkehrform ausdrücklich als Untervariante.'
  ],
  'recognition_rules': [
    'Frage: Könnte hier ein Eigenname stehen?',
    'Frage: Wird stattdessen ein Titel, Beiname oder eine Umschreibung gesetzt?',
    'Frage: Individualisiert der Ausdruck den Referenten eindeutig genug?',
    'Wenn ja, liegt wahrscheinlich Antonomasie vor.'
  ],
  'evaluation_criteria': [
    'Eindeutigkeit des Referenten',
    'funktionaler Namensersatz',
    'charakterisierende Prägnanz',
    'kontextuelle Anschlussfähigkeit',
    'saubere Abgrenzung zur Periphrase und Metapher'
  ],
  'tone_fit': [
    'epideiktisch',
    'feierlich',
    'homiletisch verdichtend',
    'prophetisch',
    'charakterisierend'
  ],
  'difficulty_level': 'mittel',
  'tags': [
    'Trope',
    'Synekdoche',
    'Eigennamenersatz',
    'Periphrase',
    'Appellativ',
    'Typusname',
    'Vossianische Antonomasie'
  ],
  'editorial_notes': 'Die Hauptfigur ist die antike Antonomasie in Lausberg §§ 580–581. Die vossianische Antonomasie wird als spätere Umkehrform integriert, aber nicht als eigene Hauptfigur geführt. Einzelne griechische Tokens im Scan sind punktuell leicht unscharf; solche Stellen sind im Material markiert.'
}

analysis_md = '''# Antonomasie – Analysepaket

## 1. Hauptfigur

**Deutscher Name:** Antonomasie  
**Griechisch:** ἀντονομασία  
**Latein:** antonomasia; ältere lateinische Nebenbezeichnung: *pronominatio*

Die Hauptfigur ist eindeutig **Antonomasie**, weil die Scan-Seite ausdrücklich mit **„F’) antonomasia (§§ 580–581)”** überschrieben ist und §580 sofort die definierende Leitsatzformel bietet. Es liegen auf der Seite zwar auch Hinweise auf Periphrase und auf die spätere vossianische Umkehrform vor, aber diese stehen klar im Dienst der einen Hauptfigur und werden von Lausberg unter demselben Lemma entfaltet. Daher ist in diesem Chat nur die Antonomasie als Hauptfigur zu behandeln.

## 2. Definition und Strukturkern

**Präzise Definition:**  
Antonomasie ist die Setzung eines **Appellativs** oder einer **Periphrase** an die Stelle eines Eigennamens. Das gemeinte Individuum wird also nicht direkt mit seinem Namen genannt, sondern durch eine kennzeichnende Benennung individualisiert.

**Einfach gesagt:**  
Nicht der Name steht da, sondern ein Titel, Beiname oder eine Umschreibung, die den Namen ersetzt.

**Strukturkern:**  
Eigenname → Appellativ / Periphrase

**Rhetorische Funktion:**
- Namensersatz mit gleichzeitiger Deutung
- Charakterisierung in einem Zug
- epideiktische Aufwertung oder Abwertung
- Verdichtung von Herkunft, Tat, Würde oder Funktion
- typologische Signalbildung

## 3. Abgrenzung und Grenzfälle

### Gegenüber der Periphrase
Nicht jede Periphrase ist schon Antonomasie. Erst wenn die Umschreibung **wirklich den Eigennamen ersetzt**, liegt Antonomasie vor.

### Gegenüber dem Epitheton
Ein Epitheton kann in antonomastische Funktion übergehen, wenn es sich vom Eigennamen löst und selbständig namensersetzend gebraucht wird.

### Gegenüber der Metapher
Die Metapher überträgt bildlich; die Antonomasie ersetzt namenshaft. Eine metaphorische Benennung ist nur dann antonomastisch, wenn sie im Kontext ein konkretes Individuum identifizierbar macht.

### Gegenüber der allgemeinen Synekdoche
Lausberg bestimmt die Antonomasie ausdrücklich als **Synekdoche für den Eigennamen**: im antiken Grundmodell entspricht dem *genus pro specie* der Synekdoche in der Antonomasie eine *species pro individuo*.

### Wichtige Grenzfälle
- bloße Apposition neben einem schon genannten Namen
- bloßer Ehrentitel ohne echte Namensfunktion
- bloße Periphrase ohne Individualisierung
- metaphorische Benennung ohne klaren Referenten

### Mit zu berücksichtigende Unterformen
1. **Appellativische Antonomasie**
2. **Periphrastische Antonomasie**
3. **Vossianische Antonomasie** als spätere Umkehrform: Eigenname → Appellativ / Typuswort

## 4. Relevante Lausberg-Stellen mit Übersetzung

### Leitdefinition (§580)

Original (Deutsch/Griechisch):
Die Antonomasie ist die Setzung eines Appellativs (λέξις) oder einer Periphrase (φράσις) an die Stelle eines Eigennamens.

Deutsch:
Die Definition ist bereits deutsch formuliert; die griechischen Zusätze benennen die beiden möglichen Ersatzformen: „Ausdruck“ und „Umschreibung“.

### Tryphon – Grunddefinition

Original (Griechisch):
ἀντονομασία ἐστὶ λέξις ἢ φράσις διὰ συνωνύμων ὀνομάτων τὸ κύριον παριστῶσα.

Deutsch:
Antonomasie ist ein Ausdruck oder eine Umschreibung, die mittels synonymischer Benennungen den Eigennamen vergegenwärtigt.

Funktionsnotiz:
Die Grundbestimmung zeigt, dass der Eigenname semantisch präsent bleibt, obwohl er sprachlich ersetzt ist.

### Griechische Kurzbeispiele aus Tryphon

Original (Griechisch):
Φοῖβος

Deutsch:
Phoibos.

Funktionsnotiz:
Ersatzname für Apollo.

Original (Griechisch):
Ἐννοσίγαιος

Deutsch:
der Erderschütterer.

Funktionsnotiz:
Ersatzname für Poseidon.

Original (Griechisch):
γλαυκῶπις [unsicher im Scan minimal angegriffen]

Deutsch:
die helläugige / eulenäugige.

Funktionsnotiz:
Ersatzname für Athene.

Original (Griechisch):
Λητοῦς καὶ Διὸς υἱός

Deutsch:
Sohn der Leto und des Zeus.

Funktionsnotiz:
Periphrastischer Ersatzname für Apollo.

### Rhetorica ad Herennium – lateinische Nebenbezeichnung

Original (Latein):
pronominatio est, quae sicuti cognomine quodam extraneo demonstrat id quod suo nomine non potest appellari

Deutsch:
Pronominatio ist eine Figur, die gleichsam durch einen fremden Beinamen das bezeichnet, was mit seinem eigenen Namen nicht so benannt werden kann.

Funktionsnotiz:
Wichtig als älterer lateinischer Terminus und als funktionale Deutung der Figur.

### Quintilian – Kurzdefinition

Original (Latein):
antonomasia, quae aliquid pro nomine ponit

Deutsch:
Antonomasie ist eine Figur, die etwas an die Stelle des Namens setzt.

Funktionsnotiz:
Maximal verdichtete klassische Definition.

### Quintilian – poetische Ersatznamen

Original (Latein):
Tydides

Deutsch:
der Sohn des Tydeus.

Funktionsnotiz:
Antonomastischer Ersatzname für Diomedes.

Original (Latein):
Pelides

Deutsch:
der Sohn des Peleus.

Funktionsnotiz:
Antonomastischer Ersatzname für Achill.

Original (Latein):
divum pater atque hominum rex

Deutsch:
Vater der Götter und König der Menschen.

Funktionsnotiz:
Periphrastischer Ersatzname für Jupiter.

### Rhetorische Ersatznamen bei Quintilian

Original (Latein):
eversorem Carthaginis et Numantiae

Deutsch:
den Zerstörer Karthagos und Numantias.

Funktionsnotiz:
Rhetorischer Ersatzname für Scipio.

Original (Latein):
Romanae eloquentiae principem

Deutsch:
den Fürsten der römischen Beredsamkeit.

Funktionsnotiz:
Rhetorischer Ersatzname für Cicero.

### Isidor – Spätantike Definition und Beispiel

Original (Latein):
antonomasia est pro nomine, id est vice nominis posita

Deutsch:
Antonomasie ist ein Ausdruck anstelle des Namens, das heißt an die Stelle des Namens gesetzt.

Funktionsnotiz:
Spätantike Wiederaufnahme der Grunddefinition.

Original (Latein):
Maia genitus

Deutsch:
der von Maia Geborene.

Funktionsnotiz:
Periphrastischer Ersatzname für Merkur.

### Weitere Lausberg-Beispiele auf der Seite

Original (Latein):
Romanus sedendo vincit

Deutsch:
Der Römer siegt durch Sitzen / Zuwarten.

Funktionsnotiz:
Romanus steht antonomastisch für Q. Fabius Maximus.

Original (Latein):
Troius heros

Deutsch:
der trojanische Held.

Funktionsnotiz:
Antonomastische Charakterisierung.

Original (Latein):
qui autem tradidit eum

Deutsch:
der aber ihn verriet.

Funktionsnotiz:
Periphrastischer Ersatzname für Judas.

Original (Latein):
discipulum quem diligebat [Iesus]

Deutsch:
den Jünger, den [Jesus] liebte.

Funktionsnotiz:
Periphrastischer Ersatzname für Johannes; die Ergänzung [Iesus] ist im Scan kontextuell erschlossen.

Original (Französisch):
Neptune le protège, et ce dieu tutélaire ne sera pas en vain imploré.

Deutsch:
Neptun schützt ihn, und dieser Schutzgott wird nicht vergeblich angerufen werden.

Funktionsnotiz:
Steigernde Periphrase statt bloßer Namenswiederholung.

Original (Französisch):
la mer qui vit tomber Icare

Deutsch:
das Meer, das Ikarus stürzen sah.

Funktionsnotiz:
Mythologisch verdichtete Ortscharakterisierung.

Original (Französisch):
cestuy-là qui conquit la toison

Deutsch:
derjenige, der das Vlies gewann.

Funktionsnotiz:
Periphrastischer Ersatzname für Jason.

Original (Englisch):
the mother of mankind

Deutsch:
die Mutter des Menschengeschlechts.

Funktionsnotiz:
Periphrastischer Ersatzname für Eva.

Original (Englisch):
one greater Man

Deutsch:
ein größerer Mensch / der größere Mensch.

Funktionsnotiz:
Christologische Antonomasie.

### §581 – Systematische Einordnung und Vossianische Umkehrform

Original (Deutsch):
Die Antonomasie ist eine Synekdoche für den Eigennamen: dem genus pro specie der Synekdoche entspricht in der Antonomasie eine species pro individuo.

Deutsch:
Die Figur ist synekdochisch gebaut: eine charakterisierende Einzelbenennung tritt für das konkrete Individuum ein.

Original (Griechisch):
πολλὴν Ἀφροδίτην τῷ λόγῳ περιτιθέναι

Deutsch:
der Rede viel Aphrodite / viel Anmut beilegen.

Funktionsnotiz:
Beispiel der vossianischen Antonomasie: Eigenname als Typuswort für eine Eigenschaft.

## 5. JSON-Datei

Die strukturierte JSON-Datei liegt separat bei als **02_Antonomasie_JSON.json**.

## 6. Einleger-Bausteine

### Titelblock
Antonomasie  
Griechisch: ἀντονομασία  
Latein: antonomasia / pronominatio

### Formel
Eigenname → Appellativ / Periphrase

### Kernwirkung
- Benennung und Deutung fallen zusammen
- Würde, Tat oder Charakter treten sofort hervor
- epideiktische Verdichtung
- typologische Aufladung

### Präzise Definition
Antonomasie ersetzt den Eigennamen durch einen kennzeichnenden Ausdruck. Dadurch wird ein Individuum nicht bloß bezeichnet, sondern zugleich interpretiert.

### Lausberg-nahe Erklärung
Nicht der Name selbst wird genannt, sondern das, was an der Person oder Sache besonders hervortritt: Herkunft, Tat, Würde, Beziehung oder Eigenschaft. So wird der Eigenname rhetorisch aufgeladen ersetzt.

### Klassische Originale
- Pelides = Achill
- Tydides = Diomedes
- divum pater atque hominum rex = Jupiter
- Maia genitus = Merkur
- the mother of mankind = Eva
- one greater Man = Christus

### Predigtnahe und moderne Beispiele
- Der Friedefürst tritt in unsere Unruhe.
- Der gute Hirte sucht verlorene Schafe.
- Der Menschenfischer zerbricht am Feuer und wird am Ufer neu gerufen.
- Die Eiserne Lady prägte ein Jahrzehnt.

### Fundus Benjamin / geschärfte deutsche Beispiele
- Sohn Gottes. Friedefürst. – Jesus.
- Der Mann der Schmerzen trägt nicht nur Wunden, sondern unsere Schuld.
- Der gute Hirte kennt nicht nur die Herde, sondern ruft dich beim Namen.

### Neuschöpfungen / predigtfähige Beispiele
- Der Friedefürst betritt nicht den Palast, sondern unsere zerrissenen Häuser.
- Der gute Hirte sucht nicht starke Tiere, sondern verlorene Schafe.
- Der Löwe aus Juda siegt nicht mit Gebrüll, sondern durch das Opfer.
- Der Menschenfischer zieht nicht Netze voll Fische, sondern Herzen aus der Tiefe.
- Der Bräutigam kommt nicht zu Satten, sondern zu Wartenden.

### Grenzfälle / Kombinationen
- Nicht jede Periphrase ist Antonomasie.
- Nicht jedes Epitheton ist Antonomasie.
- Metapher + Individualisierung kann antonomastisch werden, muss es aber nicht.
- Die vossianische Antonomasie ist eine Untervariante, keine neue Hauptfigur.

### Bauanleitung
1. Wähle die Person oder Sache.
2. Bestimme die prägnanteste Eigenschaft, Tat oder Beziehung.
3. Forme daraus einen Titel, Beinamen oder eine Periphrase.
4. Prüfe, ob der Referent ohne Eigennamen eindeutig bleibt.
5. Prüfe, ob der Ausdruck nicht bloß schmückt, sondern wirklich ersetzt.
6. Setze die Figur dort ein, wo Benennung und Auslegung in einem Schlag zusammenfallen sollen.

### Sprechschule / Merksätze
- Antonomasie nennt nicht nur; sie deutet.
- Wo der Name weicht, tritt das Profil hervor.
- Ein guter antonomastischer Ausdruck ersetzt den Namen und erklärt ihn zugleich.
- Nicht jeder Titel ist Antonomasie; entscheidend ist der echte Namensersatz.

## 7. Offene Punkte / Rückfrage an mich

Kein sachlicher Hauptzweifel an der Figurenbestimmung. Offen bleibt nur die redaktionelle Projektentscheidung, ob die **vossianische Antonomasie** künftig immer als Unterabschnitt der Antonomasie mitgeführt oder in einer späteren Datei als separate Untervarianten-Notiz ausgelagert werden soll.
'''

quellen_md = '''# Antonomasie – Quellen- und Transkriptionsnotizen

## Primärquelle
- Heinrich Lausberg, *Handbuch der literarischen Rhetorik*, Scan **Antonomasie.pdf**, Seitenbereich mit §§ 580–581.

## Projektinterne Stilreferenzen
- *MEISTERREDNER_Docx_Build_Instruktion.md*
- *MEISTERREDNER_Docx_Build_Profile.json*
- *Einleger_Stilmittel_Tiefgang_A4.docx*
- *Synonymia_Einleger_Tiefgang_A4_final2.docx*

## Sichtbare Leitstellen aus dem Scan

### §580
Die Antonomasie ist die Setzung eines Appellativs (λέξις) oder einer Periphrase (φράσις) an die Stelle eines Eigennamens.

### Tryphon
ἀντονομασία ἐστὶ λέξις ἢ φράσις διὰ συνωνύμων ὀνομάτων τὸ κύριον παριστῶσα.

Anschließende Beispiele im Scan:
- Φοῖβος
- Ἐννοσίγαιος
- γλαυκῶπις [leicht unsicher an den Rändern des Scans]
- Λητοῦς καὶ Διὸς υἱός

### Lateinische Klarstellen im Scan
- pronominatio est, quae sicuti cognomine quodam extraneo demonstrat id quod suo nomine non potest appellari
- antonomasia, quae aliquid pro nomine ponit
- antonomasia est pro nomine, id est vice nominis posita

### Weitere Beispiele auf der rechten Scanseite
- Romanus (= Q. Fabius Maximus) sedendo vincit
- Troius heros
- qui autem tradidit eum
- discipulum quem diligebat [Iesus]
- Neptune le protège, et ce dieu tutélaire ne sera pas en vain imploré
- la mer qui vit tomber Icare
- cestuy-là qui conquit la toison
- the mother of mankind
- one greater Man

### §581
Die Antonomasie ist eine Synekdoche für den Eigennamen.

Wichtige Folgeform:
- vossianische Antonomasie = Setzung eines Eigennamens für ein Appellativ / Typuswort

Griechisches Beispiel:
- πολλὴν Ἀφροδίτην τῷ λόγῳ περιτιθέναι

## Transkriptionsregel für diesen Pack
- Unsichere Lesungen wurden markiert.
- Übersetzungen sind bewusst möglichst wörtlich gehalten.
- Englische Originalbeispiele aus Lausberg wurden erhalten, nicht eingedeutscht, aber deutsch übersetzt.
'''

einleger_md = '''# Antonomasie – Einleger-Entwurf

MEISTERREDNER | Stilmittel-Einleger | Tiefgang

## Antonomasie

**Griechisch:** ἀντονομασία  
**Latein:** antonomasia / pronominatio  
**Formel:** Eigenname → Appellativ / Periphrase  
**Kernwirkung:** Verdichtung · Charakterisierung · epideiktische Schärfung · typologische Aufladung

Antonomasie ersetzt den Eigennamen durch einen kennzeichnenden Ausdruck. Nicht bloß der Name wird weggelassen; vielmehr tritt das hervor, was an der Person oder Sache besonders markant ist: Würde, Tat, Herkunft, Beziehung, Charakter.

## Lausberg-nah | Definition, Präzisierung und Wirkung

Die Antonomasie ist die Setzung eines Appellativs oder einer Periphrase an die Stelle eines Eigennamens. Sie ist ein Spezialfall der Synekdoche für Eigennamen. Das Individuum wird über eine hervorstechende Eigenschaft oder Umschreibung bezeichnet; so fallen Benennung und Deutung zusammen.

**Merksatz:** Antonomasie nennt nicht nur – sie deutet.

## Originale | Definitionen und klassische Muster

**Original (Griechisch):**
ἀντονομασία ἐστὶ λέξις ἢ φράσις διὰ συνωνύμων ὀνομάτων τὸ κύριον παριστῶσα.

**Deutsch:**
Antonomasie ist ein Ausdruck oder eine Umschreibung, die mittels synonymischer Benennungen den Eigennamen vergegenwärtigt.

**Original (Latein):**
antonomasia, quae aliquid pro nomine ponit.

**Deutsch:**
Antonomasie ist eine Figur, die etwas an die Stelle des Namens setzt.

**Original (Latein):**
antonomasia est pro nomine, id est vice nominis posita.

**Deutsch:**
Antonomasie ist ein Ausdruck anstelle des Namens, das heißt an die Stelle des Namens gesetzt.

## Klassische Beispiele | Original, Übersetzung, Markierung

**Original (Griechisch):**
Φοῖβος

**Deutsch:**
Phoibos.

**Funktionsnotiz:**
Ersatzname für Apollo.

**Original (Griechisch):**
Ἐννοσίγαιος

**Deutsch:**
der Erderschütterer.

**Funktionsnotiz:**
Ersatzname für Poseidon.

**Original (Griechisch):**
Λητοῦς καὶ Διὸς υἱός

**Deutsch:**
Sohn der Leto und des Zeus.

**Funktionsnotiz:**
Periphrastischer Ersatzname für Apollo.

**Original (Latein):**
Tydides

**Deutsch:**
der Sohn des Tydeus.

**Funktionsnotiz:**
Ersatzname für Diomedes.

**Original (Latein):**
Pelides

**Deutsch:**
der Sohn des Peleus.

**Funktionsnotiz:**
Ersatzname für Achill.

**Original (Latein):**
divum pater atque hominum rex

**Deutsch:**
Vater der Götter und König der Menschen.

**Funktionsnotiz:**
Periphrastischer Ersatzname für Jupiter.

**Original (Latein):**
Maia genitus

**Deutsch:**
der von Maia Geborene.

**Funktionsnotiz:**
Periphrastischer Ersatzname für Merkur.

**Original (Französisch):**
Neptune le protège, et ce dieu tutélaire ne sera pas en vain imploré.

**Deutsch:**
Neptun schützt ihn, und dieser Schutzgott wird nicht vergeblich angerufen werden.

**Funktionsnotiz:**
Steigernde Periphrase statt bloßer Namenswiederholung.

**Original (Französisch):**
cestuy-là qui conquit la toison

**Deutsch:**
derjenige, der das Vlies gewann.

**Funktionsnotiz:**
Periphrastischer Ersatzname für Jason.

**Original (Englisch):**
the mother of mankind

**Deutsch:**
die Mutter des Menschengeschlechts.

**Funktionsnotiz:**
Periphrastischer Ersatzname für Eva.

**Original (Englisch):**
one greater Man

**Deutsch:**
ein größerer Mensch / der größere Mensch.

**Funktionsnotiz:**
Christologische Antonomasie.

## Fundus Benjamin | geschärfte deutsche Beispiele

- Sohn Gottes. Friedefürst. – Jesus.
- Der Mann der Schmerzen trägt nicht nur Wunden, sondern unsere Schuld.
- Der gute Hirte kennt nicht nur die Herde, sondern ruft dich beim Namen.
- Der Menschenfischer zerbricht am Feuer und wird am Ufer neu gerufen.

## Neuschöpfungen | predigtfähig und klar markiert

- Der Friedefürst betritt nicht den Palast, sondern unsere zerrissenen Häuser.
- Der gute Hirte sucht nicht starke Tiere, sondern verlorene Schafe.
- Der Löwe aus Juda siegt nicht mit Gebrüll, sondern durch das Opfer.
- Der Bräutigam kommt nicht zu Satten, sondern zu Wartenden.
- Der Mann der Schmerzen trägt nicht nur den Schmerz, sondern den Träger des Schmerzes selbst.

## Grenzfälle | was oft verwechselt wird

- Nicht jede Periphrase ist Antonomasie.
- Nicht jedes Epitheton ist Antonomasie.
- Metaphorische Benennung ist erst dann antonomastisch, wenn sie ein konkretes Individuum eindeutig markiert.
- Die vossianische Antonomasie ist eine spätere Umkehrform innerhalb des Lemmas.

## Vossianische Antonomasie | Untervariante

Hier wird nicht ein Eigenname ersetzt, sondern umgekehrt ein Eigenname als Typuswort für eine Eigenschaft oder Rolle gebraucht.

**Original (Griechisch):**
πολλὴν Ἀφροδίτην τῷ λόγῳ περιτιθέναι

**Deutsch:**
der Rede viel Aphrodite / viel Anmut beilegen.

**Funktionsnotiz:**
Aphrodite steht typologisch für Anmut.

## Bauanleitung | wie du Antonomasie baust

1. Wähle die Person oder Sache.
2. Frage: Was ist an ihr im gegebenen Kontext das Prägnanteste?
3. Forme daraus einen Titel, Beinamen oder eine Umschreibung.
4. Prüfe, ob der Referent ohne Eigennamen eindeutig bleibt.
5. Prüfe, ob der Ausdruck wirklich ersetzt und nicht bloß schmückt.
6. Nutze die Figur dort, wo Benennung und Deutung in einem Schlag zusammenfallen sollen.

## Warnung

Eine unklare Antonomasie bleibt nur dekorative Umschreibung. Ohne eindeutigen Referenten verliert die Figur ihre Schärfe.

## Merksätze

- Antonomasie nennt nicht nur; sie deutet.
- Wo der Name weicht, tritt das Profil hervor.
- Ein guter antonomastischer Ausdruck ersetzt den Namen und erklärt ihn zugleich.
- Nicht jeder Titel ist Antonomasie; entscheidend ist der echte Namensersatz.
'''

readme_text = '''Antonomasie – Output-Paket

Enthaltene Dateien:
1. 01_Antonomasie_Analyse.md
   Analyse nach der Projektgliederung: Hauptfigur, Abgrenzung, Lausberg-Stellen, offene Punkte.

2. 02_Antonomasie_JSON.json
   Maschinenlesbare Korpusdatei nach dem vorgegebenen Schema.

3. 03_Antonomasie_Einleger_Entwurf.docx
   Editierbarer Word-Einleger im Stil der vorhandenen MEISTERREDNER-Dateien.

4. 04_Antonomasie_Einleger_Entwurf.md
   Textfassung des Einleger-Entwurfs.

5. 05_Quellen_und_Transkription.md
   Primärquellen-Hinweise und philologische Transkriptionsnotizen.

6. build_antonomasie_pack.py
   Erzeugungsskript für diesen Pack.

Hinweis:
Der Einleger führt die vossianische Antonomasie als Untervariante innerhalb derselben Hauptfigur.
'''

# write text artifacts
MD_ANALYSE.write_text(analysis_md, encoding='utf-8')
JSON_PATH.write_text(json.dumps(figure_json, ensure_ascii=False, indent=2), encoding='utf-8')
MD_EINLEGER.write_text(einleger_md, encoding='utf-8')
QUELLEN.write_text(quellen_md, encoding='utf-8')
README.write_text(readme_text, encoding='utf-8')

# DOCX helpers

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, sz='8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:' + edge
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), sz)
        element.set(qn('w:color'), color)


def set_table_borders(table, color=BORDER, sz='8'):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color=color, sz=sz)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)


def style_run(run, size=12, bold=False, italic=False, color=TEXT, font='Aptos'):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text='', style=None, size=12, bold=False, italic=False, color=TEXT, align=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_label_value(p, label, value, label_color=MAROON, value_color=TEXT, size=11.5):
    r1 = p.add_run(label)
    style_run(r1, size=size, bold=True, color=label_color)
    r2 = p.add_run(value)
    style_run(r2, size=size, color=value_color)


def add_box(doc, title, paragraphs, title_color=GOLD, fill=LIGHT2):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Cm(16.5)
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=BORDER, sz='8')
    cell.width = Cm(16.5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    style_run(r, size=13, bold=True, color=title_color)
    for para in paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        if isinstance(para, tuple) and para[0] == 'merksatz':
            add_label_value(p, 'Merksatz: ', para[1], label_color=MAROON, value_color=MAROON, size=11.5)
        else:
            r = p.add_run(para)
            style_run(r, size=11.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=15, bold=True, color=GOLD)
    return p


def add_example(doc, original_lang, original, deutsch, note=None, source=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(2)
    add_label_value(p, f'Original ({original_lang}):\n', '')
    r = p.add_run(original)
    style_run(r, size=11.5, italic=True, color=MAROON)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.2)
    p2.paragraph_format.space_after = Pt(2)
    add_label_value(p2, 'Deutsch:\n', '')
    r2 = p2.add_run(deutsch)
    style_run(r2, size=11.5)

    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.2)
        p3.paragraph_format.space_after = Pt(2)
        add_label_value(p3, 'Funktionsnotiz:\n', '')
        r3 = p3.add_run(note)
        style_run(r3, size=11.2)

    if source:
        p4 = doc.add_paragraph()
        p4.paragraph_format.left_indent = Cm(0.2)
        p4.paragraph_format.space_after = Pt(8)
        add_label_value(p4, 'Quelle:\n', '')
        r4 = p4.add_run(source)
        style_run(r4, size=10.5, color=GREY)


# create document

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rh = header.add_run('MEISTERREDNER | Stilmittel-Einleger | Tiefgang')
style_run(rh, size=9.5, color=GREY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run('MEISTERREDNER | Stilmittel-Einleger Tiefgang')
style_run(rf, size=9.5, color=GREY)

# title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(8)
r = p.add_run('Antonomasie')
style_run(r, size=26, bold=True, color=MAROON)

# matrix table
matrix = doc.add_table(rows=2, cols=4)
matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
matrix.autofit = False
widths = [Cm(4.0), Cm(4.0), Cm(4.6), Cm(4.7)]
headers = ['Griechisch', 'Latein', 'Formel', 'Kernwirkung']
values = [
    'ἀντονομασία',
    'antonomasia\npronominatio',
    'Eigenname →\nAppellativ /\nPeriphrase',
    'Verdichtung ·\nCharakterisierung ·\nepideiktische\nSchärfung ·\ntypologische\nAufladung'
]
for i, w in enumerate(widths):
    matrix.columns[i].width = w
    matrix.cell(0, i).width = w
    matrix.cell(1, i).width = w
for i, head in enumerate(headers):
    c = matrix.cell(0, i)
    set_cell_shading(c, LIGHT)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(head)
    style_run(r, size=13, bold=True, color=MAROON)
for i, val in enumerate(values):
    c = matrix.cell(1, i)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, line in enumerate(val.split('\n')):
        rr = p.add_run(line)
        style_run(rr, size=12, color=TEXT)
        if idx < len(val.split('\n')) - 1:
            rr.add_break()
set_table_borders(matrix)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# intro
p = add_para(doc, '', size=12)
add_label_value(p, '', 'Ersetzung eines Eigennamens durch einen Titel, Beinamen oder eine kennzeichnende Umschreibung. ', size=12.5)
r = p.add_run('Nicht der Name selbst steht da; vielmehr tritt das hervor, was die Person oder Sache in diesem Kontext besonders markiert.')
style_run(r, size=12.5, bold=True, color=TEXT)
p.paragraph_format.space_after = Pt(6)

p2 = add_para(doc, 'Typisch ist die Figur dort, wo Benennung und Deutung in einem Schlag zusammenfallen: Würde, Herkunft, Tat, Funktion oder Beziehung erscheinen sofort mit.', size=12.1)
r2 = p2.add_run(' Wo der Name weicht, tritt das Profil hervor.')
style_run(r2, size=12.1, bold=True, color=MAROON)

add_box(doc, 'Lausberg-nah | Definition, Präzisierung und Wirkung', [
    'Lausberg definiert die Antonomasie als die Setzung eines Appellativs (λέξις) oder einer Periphrase (φράσις) an die Stelle eines Eigennamens. Damit ist die Figur präzise auf den Namensersatz festgelegt.',
    'Sie ist nicht bloß schmückende Umschreibung, sondern individualisierende Benennung. Der Ausdruck muss den Eigennamen funktional ersetzen und den Referenten ausreichend kenntlich machen.',
    'Systematisch versteht Lausberg die Antonomasie als Synekdoche für den Eigennamen. In §581 notiert er außerdem die spätere vossianische Umkehrform, in der ein Eigenname typologisch als Appellativ gebraucht wird.',
    ('merksatz', 'Antonomasie nennt nicht nur – sie deutet.')
])

add_section_heading(doc, 'Originale | Definitionen und klassische Muster')
add_box(doc, '', [
    'Original (Griechisch):\nἀντονομασία ἐστὶ λέξις ἢ φράσις διὰ συνωνύμων ὀνομάτων τὸ κύριον παριστῶσα.',
    'Deutsch:\nAntonomasie ist ein Ausdruck oder eine Umschreibung, die mittels synonymischer Benennungen den Eigennamen vergegenwärtigt.',
    'Original (Latein):\nantonomasia, quae aliquid pro nomine ponit.',
    'Deutsch:\nAntonomasie ist eine Figur, die etwas an die Stelle des Namens setzt.',
    'Original (Latein):\nantonomasia est pro nomine, id est vice nominis posita.',
    'Deutsch:\nAntonomasie ist ein Ausdruck anstelle des Namens, das heißt an die Stelle des Namens gesetzt.',
    ('merksatz', 'Ein guter antonomastischer Ausdruck ersetzt den Namen und erklärt ihn zugleich.')
], fill=LIGHT2)

add_section_heading(doc, 'Klassische Beispiele | Original, Übersetzung, Markierung')
examples_1 = [
    ('Griechisch', 'Φοῖβος', 'Phoibos.', 'Ersatzname für Apollo.', 'Tryphon, bei Lausberg §580'),
    ('Griechisch', 'Ἐννοσίγαιος', 'der Erderschütterer.', 'Ersatzname für Poseidon.', 'Tryphon, bei Lausberg §580'),
    ('Griechisch', 'γλαυκῶπις [leicht unsicher im Scan]', 'die helläugige / eulenäugige.', 'Wahrscheinlich Ersatzname für Athene.', 'Tryphon, bei Lausberg §580'),
    ('Griechisch', 'Λητοῦς καὶ Διὸς υἱός', 'Sohn der Leto und des Zeus.', 'Periphrastischer Ersatzname für Apollo.', 'Tryphon, bei Lausberg §580'),
    ('Latein', 'Tydides', 'der Sohn des Tydeus.', 'Ersatzname für Diomedes.', 'Quintilian 8,6,29–30'),
    ('Latein', 'Pelides', 'der Sohn des Peleus.', 'Ersatzname für Achill.', 'Quintilian 8,6,29–30'),
    ('Latein', 'divum pater atque hominum rex', 'Vater der Götter und König der Menschen.', 'Periphrastischer Ersatzname für Jupiter.', 'Vergil, Aen. 1,65; bei Lausberg §580'),
]
for ex in examples_1:
    add_example(doc, *ex)

add_section_heading(doc, 'Weitere klassische Beispiele | rhetorisch, biblisch, französisch, englisch')
examples_2 = [
    ('Latein', 'Maia genitus', 'der von Maia Geborene.', 'Periphrastischer Ersatzname für Merkur.', 'Isidor 1,37,11; bei Lausberg §580'),
    ('Latein', 'Romanus sedendo vincit', 'Der Römer siegt durch Sitzen / Zuwarten.', 'Romanus steht antonomastisch für Q. Fabius Maximus.', 'Varro rust. 1,2,2; bei Lausberg §580'),
    ('Latein', 'Troius heros', 'der trojanische Held.', 'Antonomastische Charakterisierung.', 'Vergil, Aen. 6,451; bei Lausberg §580'),
    ('Latein', 'qui autem tradidit eum', 'der aber ihn verriet.', 'Periphrastischer Ersatzname für Judas.', 'Mt 26,48; bei Lausberg §580'),
    ('Latein', 'discipulum quem diligebat [Iesus]', 'den Jünger, den [Jesus] liebte.', 'Periphrastischer Ersatzname für Johannes.', 'Joh 19,26; bei Lausberg §580'),
    ('Französisch', 'Neptune le protège, et ce dieu tutélaire ne sera pas en vain imploré.', 'Neptun schützt ihn, und dieser Schutzgott wird nicht vergeblich angerufen werden.', 'Charakterisierend-steigernde Periphrase.', 'Phèdre 2,5,621; bei Lausberg §580'),
    ('Französisch', 'cestuy-là qui conquit la toison', 'derjenige, der das Vlies gewann.', 'Periphrastischer Ersatzname für Jason.', 'Du Bellay, Regrets 31; bei Lausberg §580'),
    ('Englisch', 'the mother of mankind', 'die Mutter des Menschengeschlechts.', 'Periphrastischer Ersatzname für Eva.', 'Milton, Paradise Lost 1,36; bei Lausberg §580'),
    ('Englisch', 'one greater Man', 'ein größerer Mensch / der größere Mensch.', 'Christologische Antonomasie.', 'Milton, Paradise Lost 1,4; bei Lausberg §580'),
]
for ex in examples_2:
    add_example(doc, *ex)

add_section_heading(doc, 'Fundus Benjamin | geschärfte deutsche Beispiele')
for item in [
    'Sohn Gottes. Friedefürst. – Jesus.',
    'Der Mann der Schmerzen trägt nicht nur Wunden, sondern unsere Schuld.',
    'Der gute Hirte kennt nicht nur die Herde, sondern ruft dich beim Namen.',
    'Der Menschenfischer zerbricht am Feuer und wird am Ufer neu gerufen.',
    'Der Friedefürst tritt nicht in dekorative Weihnachtsstimmung, sondern in zerrissene Häuser.',
]:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(4)
    rr = p.add_run('• ' + item)
    style_run(rr, size=11.8, color=TEXT)

add_section_heading(doc, 'Neuschöpfungen | predigtfähig, aber klar als moderne Ergänzung markiert')
for item in [
    'Der Friedefürst betritt nicht den Palast, sondern unsere zerrissenen Häuser.',
    'Der gute Hirte sucht nicht starke Tiere, sondern verlorene Schafe.',
    'Der Löwe aus Juda siegt nicht mit Gebrüll, sondern durch das Opfer.',
    'Der Bräutigam kommt nicht zu Satten, sondern zu Wartenden.',
    'Der Menschenfischer zieht nicht bloß Netze voll Fische, sondern Herzen aus der Tiefe.',
    'Der Mann der Schmerzen trägt nicht nur Schmerz, sondern den Träger des Schmerzes selbst heim.',
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(4)
    rr = p.add_run('• ' + item)
    style_run(rr, size=11.8, color=TEXT)

add_section_heading(doc, 'Kombinationsformen und Grenzfälle | sauber abgrenzen')
add_box(doc, '', [
    'Nicht jede Periphrase ist Antonomasie. Erst der funktionale Namensersatz entscheidet.',
    'Nicht jedes Epitheton ist Antonomasie. Das Beiwort muss sich vom Namen lösen und selbständig tragen.',
    'Nicht jede Metapher ist Antonomasie. Es braucht einen eindeutigen Individualbezug.',
    'Die vossianische Antonomasie ist keine neue Hauptfigur, sondern eine spätere Umkehrform innerhalb desselben Lemmas.',
], fill=LIGHT2)

add_section_heading(doc, 'Vossianische Antonomasie | Untervariante im selben Lemma')
add_example(doc, 'Griechisch', 'πολλὴν Ἀφροδίτην τῷ λόγῳ περιτιθέναι', 'der Rede viel Aphrodite / viel Anmut beilegen.', 'Eigenname wird typologisch als Appellativ für eine Eigenschaft benutzt.', 'Dion. Hal. de comp. 3; bei Lausberg §581')

add_section_heading(doc, 'Bauanleitung | wie du Antonomasie für Predigten baust')
for i, item in enumerate([
    'Wähle die Person oder Sache.',
    'Bestimme die eine Eigenschaft, Tat, Herkunft oder Beziehung, die im Kontext am stärksten hervortritt.',
    'Forme daraus einen Titel, Beinamen oder eine knappe Umschreibung.',
    'Prüfe, ob der Referent auch ohne Eigennamen eindeutig bleibt.',
    'Prüfe, ob die Benennung wirklich ersetzt und nicht nur dekorativ ergänzt.',
    'Nutze die Figur dort, wo Benennung und Deutung in einem Schlag zusammenfallen sollen.'
], start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(4)
    add_label_value(p, f'{i}. ', item, label_color=MAROON, value_color=TEXT, size=11.7)

add_section_heading(doc, 'Warnung und Merksätze')
add_box(doc, '', [
    'Eine unklare Antonomasie bleibt nur dekorative Umschreibung. Ohne eindeutigen Referenten verliert die Figur ihre Schärfe.',
    ('merksatz', 'Antonomasie nennt nicht nur – sie deutet.'),
    ('merksatz', 'Wo der Name weicht, tritt das Profil hervor.'),
    ('merksatz', 'Nicht jeder Titel ist Antonomasie; entscheidend ist der echte Namensersatz.'),
], fill=LIGHT2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
r = p.add_run('Quelle: Lausberg, Handbuch der literarischen Rhetorik, §§ 580–581 (Scan „Antonomasie.pdf“).')
style_run(r, size=9.7, color=GREY)

# Save and zip
DOCX_PATH.unlink(missing_ok=True)
doc.save(DOCX_PATH)

if ZIP_PATH.exists():
    ZIP_PATH.unlink()
with zipfile.ZipFile(ZIP_PATH, 'w', zipfile.ZIP_DEFLATED) as zf:
    for path in sorted(OUT.iterdir()):
        if path.name != ZIP_PATH.name:
            zf.write(path, arcname=path.name)

print('WROTE', DOCX_PATH)
print('WROTE', ZIP_PATH)
